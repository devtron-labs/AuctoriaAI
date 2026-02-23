"""
Service layer for draft document export (PDF and DOCX).

Responsibilities:
- Fetch DraftVersion by ID from the database.
- Generate PDF in memory using reportlab, preserving markdown structure
  (headings H1-H3, bold, italic, bullet lists, numbered lists).
- Generate DOCX in memory using python-docx, mapping markdown to native
  Word styles including inline bold/italic and list styles.
- Return (BytesIO buffer, filename) tuples for streaming to the client.

No files are written to disk; all generation is entirely in-memory.
"""

import io
import logging
import re

from sqlalchemy.orm import Session

from app.models.models import DraftVersion
from app.services.exceptions import NotFoundError

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _get_draft_or_404(db: Session, draft_id: str) -> DraftVersion:
    draft = db.query(DraftVersion).filter(DraftVersion.id == draft_id).first()
    if draft is None:
        raise NotFoundError(f"Draft {draft_id} not found")
    return draft


def _md_inline_to_rl(text: str) -> str:
    """Convert inline markdown to ReportLab XML markup for use inside Paragraph.

    Order matters: escape XML first, then convert bold+italic before bold/italic
    individually so that *** is not misread as * followed by **.
    """
    # Escape XML special characters so ReportLab does not choke
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # Bold + italic: ***text***
    text = re.sub(r"\*{3}([^*\n]+)\*{3}", r"<b><i>\1</i></b>", text)
    # Bold: **text** or __text__
    text = re.sub(r"\*{2}([^*\n]+)\*{2}", r"<b>\1</b>", text)
    text = re.sub(r"_{2}([^_\n]+)_{2}", r"<b>\1</b>", text)
    # Italic: *text* or _text_
    text = re.sub(r"\*([^*\n]+)\*", r"<i>\1</i>", text)
    text = re.sub(r"_([^_\n]+)_", r"<i>\1</i>", text)
    # Inline code
    text = re.sub(r"`([^`\n]+)`", r'<font name="Courier">\1</font>', text)
    # Links: keep the label text only
    text = re.sub(r"\[([^\]\n]+)\]\([^\)\n]+\)", r"\1", text)
    return text


# Compiled pattern for inline markdown tokens used in DOCX runs
_DOCX_INLINE = re.compile(
    r"\*{3}([^*\n]+)\*{3}"        # bold + italic  (group 1)
    r"|\*{2}([^*\n]+)\*{2}"       # bold           (group 2)
    r"|\*([^*\n]+)\*"             # italic         (group 3)
    r"|_{2}([^_\n]+)_{2}"         # bold alt       (group 4)
    r"|_([^_\n]+)_"               # italic alt     (group 5)
    r"|`([^`\n]+)`"               # inline code    (group 6)
    r"|\[([^\]\n]+)\]\([^\)\n]+\)"  # link label   (group 7)
)


def _add_docx_runs(paragraph, text: str) -> None:
    """Add text to a python-docx paragraph with inline markdown formatting applied."""
    last = 0
    for m in _DOCX_INLINE.finditer(text):
        # Plain text before this token
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])

        bold_italic, bold, italic, bold_alt, italic_alt, code, link_label = m.groups()

        if bold_italic is not None:
            run = paragraph.add_run(bold_italic)
            run.bold = True
            run.italic = True
        elif bold is not None:
            run = paragraph.add_run(bold)
            run.bold = True
        elif italic is not None:
            run = paragraph.add_run(italic)
            run.italic = True
        elif bold_alt is not None:
            run = paragraph.add_run(bold_alt)
            run.bold = True
        elif italic_alt is not None:
            run = paragraph.add_run(italic_alt)
            run.italic = True
        elif code is not None:
            run = paragraph.add_run(code)
            run.font.name = "Courier New"
        elif link_label is not None:
            paragraph.add_run(link_label)

        last = m.end()

    # Remaining plain text after the last token
    if last < len(text):
        paragraph.add_run(text[last:])


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

def generate_pdf(db: Session, draft_id: str) -> tuple[io.BytesIO, str]:
    """
    Generate a PDF document from a DraftVersion's markdown content.

    Markdown structure is preserved:
    - # / ## / ### headings → Heading1 / Heading2 / Heading3 paragraph styles
    - **bold** / *italic* / `code` → inline ReportLab XML markup
    - - item / * item bullet lines → indented bullet paragraphs
    - 1. item numbered lines → indented numbered paragraphs
    - --- horizontal rules → HRFlowable
    - Blank lines → small vertical spacers

    Args:
        db:       Active SQLAlchemy session.
        draft_id: UUID string of the DraftVersion.

    Returns:
        (buffer, filename): In-memory BytesIO ready for streaming and the
        suggested ``Content-Disposition`` filename.

    Raises:
        NotFoundError: If the DraftVersion does not exist.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

    draft = _get_draft_or_404(db, draft_id)
    logger.info("Generating PDF: draft_id=%s content_length=%d", draft_id, len(draft.content_markdown))

    base = getSampleStyleSheet()

    h1 = ParagraphStyle("DocH1", parent=base["Heading1"], spaceBefore=14, spaceAfter=6, keepWithNext=1)
    h2 = ParagraphStyle("DocH2", parent=base["Heading2"], spaceBefore=10, spaceAfter=4, keepWithNext=1)
    h3 = ParagraphStyle("DocH3", parent=base["Heading3"], spaceBefore=8,  spaceAfter=4, keepWithNext=1)
    body = ParagraphStyle("DocBody", parent=base["Normal"], leading=16, spaceAfter=8)
    bullet_style = ParagraphStyle(
        "DocBullet", parent=base["Normal"], leftIndent=24, firstLineIndent=0,
        leading=14, spaceAfter=4, bulletIndent=12,
    )
    numbered_style = ParagraphStyle(
        "DocNumbered", parent=base["Normal"], leftIndent=24, firstLineIndent=0,
        leading=14, spaceAfter=4,
    )
    code_style = ParagraphStyle(
        "DocCode", parent=base["Code"], leftIndent=12, leading=13,
        fontName="Courier", fontSize=9, spaceAfter=4, backColor=colors.Color(0.95, 0.95, 0.95),
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    story: list = []
    in_code_block = False
    code_lines: list[str] = []

    def _flush_code_block() -> None:
        if code_lines:
            for cl in code_lines:
                safe = cl.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe or "\u00a0", code_style))
            story.append(Spacer(1, 4))
        code_lines.clear()

    for raw_line in draft.content_markdown.splitlines():
        line = raw_line.rstrip()

        # Fenced code block toggle
        if line.startswith("```"):
            if in_code_block:
                _flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Headings
        if line.startswith("### "):
            story.append(Paragraph(_md_inline_to_rl(line[4:]), h3))
        elif line.startswith("## "):
            story.append(Paragraph(_md_inline_to_rl(line[3:]), h2))
        elif line.startswith("# "):
            story.append(Paragraph(_md_inline_to_rl(line[2:]), h1))

        # Bullet list
        elif re.match(r"^[-*]\s+", line):
            text = _md_inline_to_rl(re.sub(r"^[-*]\s+", "", line))
            story.append(Paragraph(f"\u2022\u00a0{text}", bullet_style))

        # Numbered list
        elif m := re.match(r"^(\d+)\.\s+(.*)", line):
            text = _md_inline_to_rl(m.group(2))
            story.append(Paragraph(f"{m.group(1)}.\u00a0{text}", numbered_style))

        # Horizontal rule
        elif re.match(r"^-{3,}$", line.strip()):
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
            story.append(Spacer(1, 4))

        # Regular paragraph
        elif line.strip():
            story.append(Paragraph(_md_inline_to_rl(line), body))

        # Blank line
        else:
            story.append(Spacer(1, 6))

    # Flush any unclosed code block
    _flush_code_block()

    if not story:
        story.append(Paragraph("(No content)", body))

    doc.build(story)
    buffer.seek(0)

    filename = f"draft-{draft_id}.pdf"
    logger.info("PDF ready: draft_id=%s filename=%s", draft_id, filename)
    return buffer, filename


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def generate_docx(db: Session, draft_id: str) -> tuple[io.BytesIO, str]:
    """
    Generate a DOCX document from a DraftVersion's markdown content.

    Markdown structure is preserved using native Word styles:
    - # / ## / ### headings → Heading 1 / Heading 2 / Heading 3
    - **bold** / *italic* / `code` → bold/italic/Courier New runs
    - - item / * item bullet lines → 'List Bullet' paragraph style
    - 1. item numbered lines      → 'List Number' paragraph style
    - Fenced code blocks          → Courier New body paragraphs
    - Blank lines                 → skipped (spacing handled by paragraph styles)

    Args:
        db:       Active SQLAlchemy session.
        draft_id: UUID string of the DraftVersion.

    Returns:
        (buffer, filename): In-memory BytesIO ready for streaming and the
        suggested ``Content-Disposition`` filename.

    Raises:
        NotFoundError: If the DraftVersion does not exist.
    """
    from docx import Document
    from docx.shared import Pt

    draft = _get_draft_or_404(db, draft_id)
    logger.info("Generating DOCX: draft_id=%s content_length=%d", draft_id, len(draft.content_markdown))

    doc = Document()

    # Give body paragraphs a little breathing room
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)

    in_code_block = False

    for raw_line in draft.content_markdown.splitlines():
        line = raw_line.rstrip()

        # Fenced code block toggle
        if line.startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            if line:
                run = p.add_run(line)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            continue

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        # Bullet list
        elif re.match(r"^[-*]\s+", line):
            text = re.sub(r"^[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_docx_runs(p, text)

        # Numbered list
        elif m := re.match(r"^\d+\.\s+(.*)", line):
            p = doc.add_paragraph(style="List Number")
            _add_docx_runs(p, m.group(1))

        # Horizontal rule — add a thin separator paragraph
        elif re.match(r"^-{3,}$", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        # Regular body paragraph with inline formatting
        elif line.strip():
            p = doc.add_paragraph()
            _add_docx_runs(p, line)

        # Blank lines are intentionally skipped; paragraph spacing handles gaps

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"draft-{draft_id}.docx"
    logger.info("DOCX ready: draft_id=%s filename=%s", draft_id, filename)
    return buffer, filename
